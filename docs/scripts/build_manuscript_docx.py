#!/usr/bin/env python3
"""Render the working manuscript into the .docx the journal asks for.

`docs/manuscript_draft.md` is the full working version: 14 figures, 13 tables,
separate Results and Discussion, plus internal working notes. «Почвоведение»
allows at most 5 tables and 5 figures, wants Results and Discussion as one
section, and obviously must not receive the notes-to-self. So this script does
not simply convert the markdown — it produces a journal-shaped extract:

  * the five figures and five tables that carry the argument stay in the main
    file and are renumbered 1..5;
  * everything else moves to a supplementary file, numbered S1.., and every
    in-text reference is rewritten to match;
  * the working-notes section is dropped.

The markdown stays the source of truth and is not modified.

Usage:
    python3 docs/scripts/build_manuscript_docx.py \
        --manuscript docs/manuscript_draft.md \
        --figures docs/figures_print --output docs/submission
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = 'Times New Roman'
BODY_SIZE = Pt(12)
TEXT_WIDTH_CM = 17.0

# --- what goes where -------------------------------------------------------
# (old number in the markdown, image stem, new number). Order is the order the
# reader meets them in the text.
MAIN_FIGURES = [
    (2, 'figR4_effort', 1),
    (4, 'figR2_categories', 2),
    (6, 'figR3_distribution_shape', 3),
    (8, 'figR5_zonal_sweep', 4),
    (10, 'figR6_depth_sweep', 5),
]
SUPP_FIGURES = [
    (1, None, 1), (3, 'figR8_density_grid', 2), (5, None, 3), (7, None, 4),
    (9, None, 5), (11, None, 6), (12, 'figR7_correlations', 7), (13, None, 8),
    (14, None, 9), (15, 'figR9_soil_type', 10),
]
MAIN_TABLES = [(1, 1), (3, 2), (4, 3), (7, 4), (13, 5)]
SUPP_TABLES = [(2, 1), (5, 2), (6, 3), (8, 4), (9, 5), (10, 6), (11, 7),
               (12, 8), (14, 9)]

FIGURE_MAP = {old: str(new) for old, _, new in MAIN_FIGURES}
FIGURE_MAP.update({old: f'S{new}' for old, _, new in SUPP_FIGURES})
TABLE_MAP = {old: str(new) for old, new in MAIN_TABLES}
TABLE_MAP.update({old: f'S{new}' for old, new in SUPP_TABLES})

# Sections of the markdown that must not reach the editor, and sections that
# are folded into «Результаты и обсуждение».
RENAME = {'ФИНАНСИРОВАНИЕ': 'ИСТОЧНИКИ ФИНАНСИРОВАНИЯ'}

DROP_SECTIONS = {'ОТКРЫТЫЕ ВОПРОСЫ И НЕОБХОДИМЫЕ ДЕЙСТВИЯ АВТОРОВ',
                 'ТАБЛИЦЫ', 'ПОДПИСИ К РИСУНКАМ', 'АННОТАЦИЯ'}
MERGE_INTO_RESULTS = ['РЕЗУЛЬТАТЫ', 'ОБСУЖДЕНИЕ', 'БАЗА ДАННЫХ И ГЕОПОРТАЛ',
                      'ОГРАНИЧЕНИЯ ИССЛЕДОВАНИЯ',
                      'ВОЗМОЖНОСТИ ПОВТОРНОГО ИСПОЛЬЗОВАНИЯ ДАННЫХ']


# --- docx plumbing ---------------------------------------------------------

def setup(document: Document, landscape_first: bool = False) -> None:
    style = document.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), BODY_FONT)
    fmt = style.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(1.0)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for section in document.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def para(document, text='', *, bold=False, italic=False, align=None,
         indent=None, space_before=0, space_after=0, size=None):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(indent if indent is not None else 1.0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_runs(p, text, bold=bold, italic=italic, size=size)
    return p


INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|<[^>\s]+>)', re.S)


def add_runs(paragraph, text, *, bold=False, italic=False, size=None):
    """Turn the markdown inline marks into runs. Code spans become italics —
    the journal has no monospace convention and file names read as titles."""
    # Unbalanced code marks appear where a span wraps across source lines;
    # dropping the marks is right, the file name still reads as a file name.
    text = text.replace('``', '')
    if text.count('`') % 2:
        text = text.replace('`', '')
    text = re.sub(r'`([^`]*)`', r'\1', text) if text.count('`') else text
    for piece in INLINE.split(text):
        if not piece:
            continue
        b, i = bold, italic
        if piece.startswith('**') and piece.endswith('**'):
            piece, b = piece[2:-2], True
        elif piece.startswith('*') and piece.endswith('*') and len(piece) > 2:
            piece, i = piece[1:-1], True
        elif piece.startswith('`') and piece.endswith('`'):
            piece, i = piece[1:-1], True
        elif piece.startswith('<') and piece.endswith('>'):
            piece = piece[1:-1]
        run = paragraph.run_style = paragraph.add_run(piece)
        run.bold, run.italic = b, i
        run.font.name = BODY_FONT
        if size is not None:
            run.font.size = size


def heading(document, text, level=1):
    p = para(document, text, bold=True, indent=0,
             space_before=12 if level == 1 else 10, space_after=6,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    if level == 2:
        for run in p.runs:
            run.italic = True
    return p


def set_borders(table):
    """The journal asks for ruled columns, which python-docx will not do for
    us — the borders have to be written into the table properties directly."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tbl_pr.append(borders)


# --- markdown parsing ------------------------------------------------------

def remap_references(text: str) -> str:
    """Rewrite «рис. 8» / «табл. 13» to the numbering of the split documents."""
    def fig(match):
        nums = [int(n) for n in re.findall(r'\d+', match.group(0))]
        return 'рис. ' + ', '.join(FIGURE_MAP.get(n, str(n)) for n in nums)

    def tab(match):
        nums = [int(n) for n in re.findall(r'\d+', match.group(0))]
        return 'табл. ' + ', '.join(TABLE_MAP.get(n, str(n)) for n in nums)

    text = re.sub(r'рис\. \d+(?:, \d+)*', fig, text)
    text = re.sub(r'табл\. \d+(?:, \d+)*', tab, text)
    return text.replace('табл. П1', 'табл. S9')


def split_sections(markdown: str) -> list[tuple[int, str, str]]:
    """-> [(level, title, body)] for every ## and ### heading."""
    parts, current = [], None
    buffer: list[str] = []
    for line in markdown.split('\n'):
        m = re.match(r'^(#{2,3}) (.+)$', line)
        if m:
            if current:
                parts.append((*current, '\n'.join(buffer).strip()))
            current, buffer = (len(m.group(1)) - 1, m.group(2).strip()), []
        elif current:
            buffer.append(line)
    if current:
        parts.append((*current, '\n'.join(buffer).strip()))
    return parts


def blocks(body: str):
    """Yield ('table', rows) and ('para', text) in document order."""
    chunk: list[str] = []
    table: list[str] = []
    for raw in body.split('\n') + ['']:
        line = raw.rstrip()
        if line.startswith('|'):
            if chunk:
                yield 'para', ' '.join(chunk)
                chunk = []
            table.append(line)
            continue
        if table:
            yield 'table', table
            table = []
        if not line.strip() or line.strip() == '---':
            if chunk:
                yield 'para', ' '.join(chunk)
                chunk = []
        else:
            chunk.append(line.strip())
    if chunk:
        yield 'para', ' '.join(chunk)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if re.match(r'^\|[\s:\-|]+\|$', line):
            continue
        rows.append([c.strip() for c in line.strip().strip('|').split('|')])
    return rows


def write_table(document, rows, *, font=Pt(10)):
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = 'Table Grid'
    set_borders(table)
    for r, row in enumerate(rows):
        for c in range(width):
            cell = table.cell(r, c)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            text = row[c] if c < len(row) else ''
            # The journal forbids empty cells: an em dash states "not applicable".
            add_runs(p, text if text else '—', bold=(r == 0), size=font)
    return table


# --- document assembly -----------------------------------------------------

def first_page(document, meta: dict) -> None:
    para(document, '[ЗАПОЛНИТЬ: рубрика журнала — одна из восьми предусмотренных]',
         indent=0, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    para(document, f"УДК {meta['udk']}", indent=0,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
    para(document, meta['title'], bold=True, indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    para(document, '© 2026 г. А. М. Гафуров', indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(document,
         'Казанский (Приволжский) федеральный университет, '
         '[ЗАПОЛНИТЬ: институт, кафедра], '
         'ул. Кремлёвская, 18, Казань, 420000, Россия',
         indent=0, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(document, 'e-mail: [ЗАПОЛНИТЬ]', indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(document, 'ORCID: 0000-0002-0812-1750', indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(document, 'Поступила в редакцию [ЗАПОЛНИТЬ]', indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    for block in meta['abstract']:
        para(document, remap_references(block), space_after=6)
    para(document, '', space_after=6)
    para(document, meta['keywords'], space_after=6)


def build_main(markdown: str, figures: Path, out: Path) -> dict:
    document = Document()
    setup(document)

    title = markdown.split('\n')[0].lstrip('# ').strip()
    udk = re.search(r'\*\*УДК\*\* (\S+)', markdown).group(1)
    abstract_src = markdown.split('## АННОТАЦИЯ')[1].split('**Ключевые слова:**')[0]
    abstract = [' '.join(p.split())
                for p in abstract_src.strip().split('\n\n') if p.strip()]
    keywords = '**Ключевые слова:** ' + ' '.join(
        markdown.split('**Ключевые слова:**')[1].split('---')[0].split())

    first_page(document, {'title': title, 'udk': udk, 'abstract': abstract,
                          'keywords': keywords})

    sections = split_sections(markdown)
    stats = {'paragraphs': 0, 'chars': 0}
    in_results = False
    for level, name, body in sections:
        if level == 1 and name in DROP_SECTIONS:
            continue
        if level == 1 and name.startswith('СПИСОК ЛИТЕРАТУРЫ'):
            break
        if level == 1 and name in MERGE_INTO_RESULTS:
            if not in_results:
                heading(document, 'РЕЗУЛЬТАТЫ И ОБСУЖДЕНИЕ', 1)
                in_results = True
            if name != 'РЕЗУЛЬТАТЫ':
                heading(document, name.capitalize(), 2)
        elif level == 1:
            in_results = False
            # The journal names this section «Источники финансирования».
            heading(document, RENAME.get(name, name), 1)
        else:
            heading(document, name, 2)
        for kind, payload in blocks(body):
            if kind == 'table':
                write_table(document, parse_table(payload))
                para(document, '', space_after=6)
            else:
                text = remap_references(payload)
                if text.strip():
                    para(document, text, space_after=0)
                    stats['paragraphs'] += 1
                    stats['chars'] += len(text)

    references(document, markdown)
    tables_section(document, markdown, MAIN_TABLES, 'ТАБЛИЦЫ')
    captions_and_figures(document, markdown, figures)

    document.save(out)
    return stats


def references(document, markdown: str) -> None:
    heading(document, 'СПИСОК ЛИТЕРАТУРЫ', 1)
    para(document,
         '*Список приведён в рабочем виде: позиции, выходные данные которых '
         'подлежат установлению, помечены. Перед подачей список необходимо '
         'переупорядочить по алфавиту (сначала русскоязычные, затем '
         'иностранные работы) и заменить упоминания в тексте на номера в '
         'квадратных скобках согласно правилам журнала.*',
         italic=True, space_after=8)
    body = markdown.split('## СПИСОК ЛИТЕРАТУРЫ')[1].split('## ТАБЛИЦЫ')[0]
    for kind, payload in blocks(body):
        if kind == 'para' and payload.strip():
            para(document, payload, space_after=0)


def tables_section(document, markdown: str, wanted, title: str) -> None:
    document.add_page_break()
    heading(document, title, 1)
    src = markdown.split('## ТАБЛИЦЫ')[1].split('## ПОДПИСИ К РИСУНКАМ')[0]
    chunks = re.split(r'\n(?=\*\*Таблица )', src)
    by_number = {}
    for chunk in chunks:
        m = re.match(r'\*\*Таблица (\d+|П1)\.?\*\*', chunk.strip())
        if m:
            key = 0 if m.group(1) == 'П1' else int(m.group(1))
            by_number[key] = chunk.strip()
    for old, new in wanted:
        chunk = by_number.get(old)
        if not chunk:
            continue
        caption, rest = chunk.split('\n', 1) if '\n' in chunk else (chunk, '')
        lines = chunk.split('\n')
        caption_lines, body_lines = [], []
        for line in lines:
            (body_lines if line.startswith('|') or body_lines
             else caption_lines).append(line)
        caption = ' '.join(' '.join(caption_lines).split())
        caption = re.sub(r'^\*\*Таблица (?:\d+|П1)\.?\*\*',
                         f'**Таблица {new}.**', caption)
        para(document, remap_references(caption), indent=0, space_before=10,
             space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT)
        table_lines = [l for l in body_lines if l.startswith('|')]
        write_table(document, parse_table(table_lines))
        tail = [l for l in body_lines if l and not l.startswith('|')]
        if tail:
            para(document, remap_references(' '.join(' '.join(tail).split())),
                 indent=0, size=Pt(10), space_before=4, space_after=8)


def captions_and_figures(document, markdown: str, figures: Path) -> None:
    document.add_page_break()
    heading(document, 'ПОДПИСИ К РИСУНКАМ', 1)
    src = markdown.split('## ПОДПИСИ К РИСУНКАМ')[1]
    chunks = re.split(r'\n(?=\*\*Рис\. )', src)
    by_number = {}
    for chunk in chunks:
        m = re.match(r'\*\*Рис\. (\d+)\.\*\*', chunk.strip())
        if m:
            by_number[int(m.group(1))] = ' '.join(chunk.strip().split())
    for old, stem, new in MAIN_FIGURES:
        caption = by_number.get(old, '')
        caption = re.sub(r'\*\*Рис\. \d+\.\*\*', f'**Рис. {new}.**', caption)
        caption = re.sub(r'\s*\*\(`[^`]+`\)\*', '', caption)
        para(document, remap_references(caption), indent=0, space_after=8,
             align=WD_ALIGN_PARAGRAPH.LEFT)

    for old, stem, new in MAIN_FIGURES:
        document.add_page_break()
        image = figures / f'{stem}.png'
        if image.exists():
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run().add_picture(str(image), width=Cm(TEXT_WIDTH_CM))
        para(document, f'Рис. {new}', indent=0, space_before=6,
             align=WD_ALIGN_PARAGRAPH.CENTER)


def build_supplement(markdown: str, figures: Path, out: Path) -> None:
    document = Document()
    setup(document)
    para(document, 'ДОПОЛНИТЕЛЬНЫЕ МАТЕРИАЛЫ', bold=True, indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    para(document, markdown.split('\n')[0].lstrip('# ').strip(), bold=True,
         indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    para(document, 'А. М. Гафуров', indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(document,
         'Ниже приведены таблицы и рисунки, не вошедшие в основной текст '
         'вследствие ограничения журнала на их число. Нумерация с префиксом S '
         'соответствует ссылкам в основном тексте.', space_after=10)

    tables_section(document, markdown, [(o, f'S{n}') for o, n in SUPP_TABLES],
                   'ТАБЛИЦЫ ДОПОЛНИТЕЛЬНЫХ МАТЕРИАЛОВ')
    src = markdown.split('## ТАБЛИЦЫ')[1].split('## ПОДПИСИ К РИСУНКАМ')[0]
    census = [c for c in re.split(r'\n(?=\*\*Таблица )', src)
              if c.strip().startswith('**Таблица П1')]
    if census:
        para(document, remap_references(' '.join(census[0].split()).replace(
            '**Таблица П1** (приложение).', '**Таблица S9.**')),
            indent=0, space_before=10, space_after=8)

    document.add_page_break()
    heading(document, 'ПОДПИСИ К РИСУНКАМ ДОПОЛНИТЕЛЬНЫХ МАТЕРИАЛОВ', 1)
    caps = markdown.split('## ПОДПИСИ К РИСУНКАМ')[1]
    chunks = re.split(r'\n(?=\*\*Рис\. )', caps)
    by_number = {}
    for chunk in chunks:
        m = re.match(r'\*\*Рис\. (\d+)\.\*\*', chunk.strip())
        if m:
            by_number[int(m.group(1))] = ' '.join(chunk.strip().split())
    for old, stem, new in SUPP_FIGURES:
        caption = by_number.get(old, '')
        caption = re.sub(r'\*\*Рис\. \d+\.\*\*', f'**Рис. S{new}.**', caption)
        caption = re.sub(r'\s*\*\(`[^`]+`\)\*', '', caption)
        para(document, remap_references(caption), indent=0, space_after=8,
             align=WD_ALIGN_PARAGRAPH.LEFT)

    for old, stem, new in SUPP_FIGURES:
        if not stem:
            continue
        image = figures / f'{stem}.png'
        if not image.exists():
            continue
        document.add_page_break()
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run().add_picture(str(image), width=Cm(TEXT_WIDTH_CM))
        para(document, f'Рис. S{new}', indent=0, space_before=6,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    document.save(out)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('--manuscript', type=Path, required=True)
    parser.add_argument('--figures', type=Path, required=True)
    parser.add_argument('--output', type=Path, required=True)
    args = parser.parse_args()
    args.output.mkdir(parents=True, exist_ok=True)

    markdown = args.manuscript.read_text(encoding='utf-8')
    main_path = args.output / 'Гафуров_Почвоведение_рукопись.docx'
    supp_path = args.output / 'Гафуров_Почвоведение_доп_материалы.docx'
    stats = build_main(markdown, args.figures, main_path)
    build_supplement(markdown, args.figures, supp_path)

    figure_dir = args.output / 'рисунки'
    figure_dir.mkdir(exist_ok=True)
    for old, stem, new in MAIN_FIGURES:
        source = args.figures / f'{stem}.png'
        if source.exists():
            (figure_dir / f'Рис_{new}.png').write_bytes(source.read_bytes())

    print(f'  {main_path.name}')
    print(f'  {supp_path.name}')
    print(f'  рисунки/ — {len(list(figure_dir.glob("*.png")))} файлов')
    print(f'  текста: {stats["chars"]} знаков '
          f'(~{stats["chars"] / 1800:.1f} стр. по 1800 знаков)')


if __name__ == '__main__':
    main()
